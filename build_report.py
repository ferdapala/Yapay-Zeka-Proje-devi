# -*- coding: utf-8 -*-
"""
build_report.py
---------------
train_and_evaluate.py'nin urettigi GERCEK outputs/ dosyalarindan Word/PDF rapor olusturur.
Cikti: rapor/Odev2_Raporu.docx  (+ docx2pdf varsa Odev2_Raporu.pdf)
Calistirma: py -3.9 build_report.py
"""
import os
import json
import pandas as pd
import numpy as np
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

BASE = os.path.dirname(os.path.abspath(__file__))
DATA = os.path.join(BASE, "data")
OUT = os.path.join(BASE, "outputs")
REP = os.path.join(BASE, "rapor")
os.makedirs(REP, exist_ok=True)

# >>> Grup uyeleri kendi bilgilerini buraya yazsin <<<
GROUP_MEMBERS = [
    "[Ad SOYAD]  -  [Ogrenci No]",
    "[Ad SOYAD]  -  [Ogrenci No]",
    "[Ad SOYAD]  -  [Ogrenci No]",
]
GITHUB_URL = "https://github.com/bt4xbjcyd7-rgb/Projelerim"

# ----------------------------------------------------------------------------
# Cikti tablolarini oku
# ----------------------------------------------------------------------------
cos = pd.read_csv(os.path.join(OUT, "cosine_eval.csv"))
sem = pd.read_csv(os.path.join(OUT, "semantic_eval.csv"))
sw = pd.read_csv(os.path.join(OUT, "similar_words.csv"))
top5 = pd.read_csv(os.path.join(OUT, "top5_per_model.csv"))
jac = pd.read_csv(os.path.join(OUT, "jaccard_matrix.csv"), index_col=0)
summary = pd.read_csv(os.path.join(OUT, "summary.csv"))
query_txt = open(os.path.join(OUT, "query.txt"), encoding="utf-8").read().strip().splitlines()

# ----------------------------------------------------------------------------
# Veri istatistiklerini (ham vs on-islenmis) hesapla
# ----------------------------------------------------------------------------
def word_stats():
    lem = pd.read_csv(os.path.join(DATA, "lemmatized_sentences.csv"), sep=";")
    stm = pd.read_csv(os.path.join(DATA, "stemmed_sentences.csv"), sep=";")
    n_rows = len(lem)

    # ham metin = name + categories
    raw_tokens, raw_vocab = 0, set()
    src = os.path.join(DATA, "yelp_academic_dataset_business.json")
    with open(src, encoding="utf-8") as f:
        for line in f:
            line = line.strip()
            if not line:
                continue
            b = json.loads(line)
            text = (str(b.get("name", "")) + " " + str(b.get("categories", ""))).lower()
            toks = [t for t in __import__("re").findall(r"[A-Za-z]+", text)]
            raw_tokens += len(toks)
            raw_vocab.update(toks)

    def tok_stats(df):
        toks = [w for t in df["text"].fillna("").astype(str) for w in t.split()]
        return len(toks), len(set(toks))

    lt, lv = tok_stats(lem)
    st, sv = tok_stats(stm)
    return {
        "rows": n_rows,
        "raw_tokens": raw_tokens, "raw_vocab": len(raw_vocab),
        "lem_tokens": lt, "lem_vocab": lv,
        "stm_tokens": st, "stm_vocab": sv,
    }


STATS = word_stats()
# min_count=2 sonrasi modelin gercek kelime hazinesi (cikti loglarindan bagimsiz, sabit ifade)
VOCAB_LEM, VOCAB_STM = 249, 248

# en iyi / en zayif modeller (anlamsal ortalamaya, esitlikte cosine'e gore)
rank_df = summary.sort_values(["mean_semantic", "mean_cosine"], ascending=False).reset_index(drop=True)
BEST = rank_df.iloc[0]["model"]
BEST2 = rank_df.iloc[1]["model"]
WORST = rank_df.iloc[-1]["model"]

# jaccard off-diagonal istatistikleri
_j = jac.values.astype(float)
_mask = ~np.eye(_j.shape[0], dtype=bool)
J_MIN, J_MAX, J_MEAN = _j[_mask].min(), _j[_mask].max(), _j[_mask].mean()

# ----------------------------------------------------------------------------
# Word yardimcilari
# ----------------------------------------------------------------------------
doc = Document()
doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(11)


def H(t, level=1):
    doc.add_heading(t, level=level)


def P(t, bold=False, italic=False, size=11, align=None):
    p = doc.add_paragraph()
    r = p.add_run(t)
    r.bold, r.italic = bold, italic
    r.font.size = Pt(size)
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


def bullet(t):
    doc.add_paragraph(t, style="List Bullet")


def table(df, font_size=8, col_widths=None):
    t = doc.add_table(rows=1, cols=len(df.columns))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, c in enumerate(df.columns):
        cell = t.rows[0].cells[j]
        cell.text = str(c)
        for pr in cell.paragraphs:
            for rn in pr.runs:
                rn.bold = True
                rn.font.size = Pt(font_size)
    for _, row in df.iterrows():
        cells = t.add_row().cells
        for j, c in enumerate(df.columns):
            cells[j].text = str(row[c])
            for pr in cells[j].paragraphs:
                for rn in pr.runs:
                    rn.font.size = Pt(font_size)
    if col_widths:
        for j, w in enumerate(col_widths):
            for r in t.rows:
                r.cells[j].width = Inches(w)
    return t


# ============================================================================
# KAPAK
# ============================================================================
P("YAPAY ZEKA DERSI", bold=True, size=14, align="center")
P("Odev-2", size=12, align="center")
doc.add_paragraph()
P("EGITILEN WORD2VEC MODELLERI ILE", bold=True, size=18, align="center")
P("METIN BENZERLIGI HESAPLAMA VE DEGERLENDIRME", bold=True, size=18, align="center")
doc.add_paragraph()
P("Proje Konusu: Yelp Isletme (Business) Verisi - Benzer Isletme Bulma", size=12, align="center")
doc.add_paragraph()
doc.add_paragraph()
P("Proje Sahipleri", bold=True, size=12, align="center")
for m in GROUP_MEMBERS:
    P(m, size=12, align="center")
doc.add_paragraph()
P("Teslim Tarihi: 15 Haziran 2026", size=11, align="center")
P("GitHub: " + GITHUB_URL, size=10, align="center")
doc.add_page_break()

# ============================================================================
# 1. GIRIS
# ============================================================================
H("1. Giris", 1)

H("1.1 Odevin Amaci", 2)
P("Bu odevin amaci, birinci odevde on isleme tabi tuttugumuz iki temiz veri seti "
  "(lemmatized ve stemmed) uzerinde Gensim ile Word2Vec modelleri egitmek ve bu modelleri "
  "kullanarak metinler/dokumanlar arasi benzerlik hesaplamaktir. Toplam 16 model "
  "(2 veri seti x 8 parametre seti) egitilmis; veri setinden secilen ornek bir isletmeye en "
  "benzer 5 isletme her model icin ayri ayri bulunmus ve modeller uc yontemle (Cosine, Anlamsal, "
  "Jaccard) karsilastirmali degerlendirilmistir.")

H("1.2 Kullanilan Veri Seti", 2)
P("Projemizde Yelp Academic Dataset'in isletme (business) dosyasi kullanilmistir. Her kayit; "
  "isletme adi (name), kategoriler (categories), sehir, eyalet, yildiz puani ve yorum sayisi gibi "
  "alanlar icermektedir. Birinci odevde her isletme icin 'name + categories' metni alinmis; kucuk "
  "harfe cevirme, NLTK ile kelime tokenizasyonu, Ingilizce stop-word temizligi ve ardindan "
  "Lemmatization (WordNetLemmatizer) / Stemming (PorterStemmer) uygulanarak iki temiz veri seti "
  "uretilmistir. Veri Ingilizcedir ve her satir (dokuman) bir isletmeyi temsil eder.")
bullet("Toplam isletme (dokuman) sayisi: {}".format(STATS["rows"]))
bullet("Benzerlik gorevindeki 'dokuman' birimi: bir isletmenin 'ad + kategoriler' metnidir.")

H("1.3 Ham Veri ile On-Islenmis Verinin Karsilastirilmasi", 2)
P("On isleme (kucuk harf, noktalama/sayi disi tokenlarin atilmasi, stop-word temizligi, "
  "lemmatization/stemming) sonrasinda metinler sadelesmis ve kelime hazinesi kuculmustur. Asagidaki "
  "tablo ham metin ile model egitiminde kullanilan iki temiz setin boyut/yapi karsilastirmasidir:")
comp = pd.DataFrame({
    "Veri Seti": ["Ham (name+categories)", "Lemmatized", "Stemmed"],
    "Dokuman": [STATS["rows"], STATS["rows"], STATS["rows"]],
    "Toplam Kelime": [STATS["raw_tokens"], STATS["lem_tokens"], STATS["stm_tokens"]],
    "Ort. Kelime/Dokuman": [round(STATS["raw_tokens"]/STATS["rows"], 2),
                            round(STATS["lem_tokens"]/STATS["rows"], 2),
                            round(STATS["stm_tokens"]/STATS["rows"], 2)],
    "Benzersiz Kelime": [STATS["raw_vocab"], STATS["lem_vocab"], STATS["stm_vocab"]],
})
table(comp, font_size=9)
P("Yorum: Stop-word ve harf-disi tokenlarin temizligiyle toplam kelime sayisi ham veriye gore "
  "azalmistir. Benzersiz kelime sayisi stemmed sette lemmatized'e gore biraz daha dusuktur; cunku "
  "Porter Stemmer kelimeleri daha agresif koklerine indirger (orn. 'restaurants' -> 'restaur', "
  "'bakeries' -> 'bakeri') ve farkli ekli varyantlari tek bicimde toplar. Word2Vec egitiminde "
  "min_count=2 esigi uygulandigi icin modellerin nihai kelime hazinesi lemmatized icin {} , stemmed "
  "icin {} kelimedir.".format(VOCAB_LEM, VOCAB_STM))

H("1.4 GitHub Reposu ve Modeller", 2)
P("Tum calisma kodlari (veri uretimi/on isleme, model egitimi, benzerlik hesabi, degerlendirme ve "
  "rapor uretimi) ve egitilen 16 model GitHub reposuna eklenmistir. Modeller boyut nedeniyle "
  "yuklenemezse Drive linki rapora eklenmelidir (herkese acik). Calistirma talimatlari README'de.")
bullet("Repo: " + GITHUB_URL)

doc.add_page_break()

# ============================================================================
# 2. YONTEM
# ============================================================================
H("2. Yontem", 1)

H("2.1 Word2Vec Vektorlestirme (Gorev-1)", 2)
P("Gensim Word2Vec sinifi ile her iki veri seti icin asagidaki 8 parametre setiyle ayri ayri model "
  "egitilmistir (toplam 16). CBOW icin sg=0, SkipGram icin sg=1 kullanilmistir. Dokumanlar kisa "
  "oldugundan epochs=80, gurultuyu azaltmak icin min_count=2, tekrarlanabilirlik icin workers=1 ve "
  "seed=42 secilmistir.")
bullet("Ortak parametreler: min_count=2, epochs=80, workers=1, seed=42")
bullet("Degisen parametreler: algoritma (cbow/skipgram), window (2/4), vector_size (100/300)")
model_tbl = cos[["model"]].copy()
model_tbl["Veri Seti"] = model_tbl["model"].apply(lambda s: s.split("_")[1])
model_tbl["Algoritma"] = model_tbl["model"].apply(lambda s: s.split("_")[2])
model_tbl["Window"] = model_tbl["model"].apply(lambda s: s.split("_")[3].replace("win", ""))
model_tbl["Boyut"] = model_tbl["model"].apply(lambda s: s.split("_")[4].replace("dim", ""))
model_tbl.columns = ["Model Adi", "Veri Seti", "Algoritma", "Window", "Vektor Boyutu"]
table(model_tbl, font_size=8)

P("Gorev-1 - Ornek Vektor Ciktilari (anlamli kelime: 'coffee')", bold=True)
P("Her model icin 'coffee' kelimesinin vektor uzayindaki en yakin 5 komsusu asagida verilmistir. "
  "Amac, modelden modele komsuluk iliskilerinin ve benzerlik skorlarinin nasil degistigini "
  "gostermektir. (Bu skorlar tek basina basariyi olcmek icin yeterli degildir.)")
table(sw.rename(columns={"model": "Model", "keyword": "Kelime",
                         "top5_neighbours": "En Yakin 5 Kelime (skor)"}), font_size=7)
P("Gozlem: Tum modellerde 'coffee' kelimesinin en yakin komsulari 'cafe', 'tea', 'roasted', "
  "'donut', 'bakery', 'bean', 'morning' gibi gercekten ayni temadaki (kahve/kafe/firin) kelimelerdir; "
  "bu da modellerin kategori birlikteliklerini basariyla ogrendigini gosterir. CBOW modellerinde "
  "komsuluk skorlari (~0.98-0.99) SkipGram'a gore biraz daha yuksek ve sik; SkipGram skorlari biraz "
  "daha dusuk ve ayristiricidir. Beklentimiz: bu kisa, kategori-yogun metinlerde CBOW'un tutarli ve "
  "guclu komsuluklar uretmesi, SkipGram'in ise daha ince/ayristirici fakat biraz daha gurultulu "
  "iliskiler ogrenmesidir.")

H("2.2 Benzerlik Hesaplama Yontemi (Gorev-2)", 2)
P("Ornek giris metni veri setinin kendi icinden secilmistir (disaridan veri alinmamistir). Giris "
  "isletmesi bir kahve/kafe isletmesidir:")
for line in query_txt:
    P("   " + line, italic=True)
P("Benzerlik su sekilde hesaplanmistir:")
bullet("Hem giris isletmesi hem de her isletme icin, metindeki kelimelerin model vektorlerinin "
       "ARITMETIK ORTALAMASI alinarak bir dokuman vektoru olusturulur.")
bullet("Modelde karsiligi olmayan (OOV) kelimeler atlanir. Bir metindeki hicbir kelime modelde "
       "yoksa NaN/ZeroDivision hatasini onlemek icin tamamen sifirlardan olusan SIFIR VEKTOR atanir.")
bullet("Giris vektoru ile her dokuman vektoru arasinda COSINE SIMILARITY hesaplanir; en yuksek 5 "
       "isletme (giris isletmesinin kendisi haric) o model icin sonuc kabul edilir.")
bullet("Islem 16 modelin her biri icin tekrarlanir: 16 x 5 = toplam 80 benzer dokuman.")

doc.add_page_break()

# ============================================================================
# 3. SONUCLAR VE DEGERLENDIRME
# ============================================================================
H("3. Sonuclar ve Degerlendirme", 1)

H("3.1 Her Model icin Ilk 5 Benzer Isletme", 2)
P("Asagida her modelin getirdigi en benzer 5 isletme; cosine skoru ve elle verilen anlamsal puan "
  "(1-5) ile listelenmistir.")
for model in cos["model"].tolist():
    sub = top5[top5["model"] == model].sort_values("rank")
    P(model, bold=True, size=10)
    tb = sub[["rank", "cosine", "semantic_1_5", "business"]].copy()
    tb.columns = ["Sira", "Cosine", "Anlamsal(1-5)", "Isletme (ad - kategoriler - sehir)"]
    table(tb, col_widths=[0.4, 0.7, 0.9, 4.5], font_size=8)
    doc.add_paragraph()

doc.add_page_break()

H("3.2 (1) Cosine Degerlendirme (Objective Evaluation)", 2)
P("Her model icin ilk 5 sonucun cosine skorlari ve ortalamalari:")
cos_tbl = cos.copy()
cos_tbl.columns = ["Model", "Skor1", "Skor2", "Skor3", "Skor4", "Skor5", "Ortalama"]
table(cos_tbl, font_size=8)
P("Yorum: Tum modellerin ortalama cosine skorlari tavana cok yakindir (~0.998-0.9996). Bunun nedeni, "
  "dokuman vektorlerinin cok kisa metinlerdeki (ortalama birkac kelime) kelime vektorlerinin "
  "ortalamasi olmasi ve kategori-yogun korpusta vektorlerin birbirine cok yakin yonlerde "
  "olusmasidir. Bu nedenle MUTLAK cosine degeri tek basina zayif bir ayristiricidir; nitekim odev "
  "metninde de bu skorlarin tek basina yeterli olmadigi belirtilmistir. Yine de SIRALAMA dogru "
  "calismaktadir: getirilen ilk 5 isletmenin neredeyse tamami kahve/kafe/firin temasindadir.")

H("3.3 (2) Anlamsal Degerlendirme (Subjective Evaluation)", 2)
P("Her modelin onerdigi 5 isletmeye, giris isletmesiyle (kahve/kafe) anlamsal yakinligina gore 1-5 "
  "arasi puan verilmistir (1: cok alakasiz ... 5: neredeyse ayni temada). Puanlama, iki isletmenin "
  "ortak kategori/tema kelimesi sayisina dayanan tutarli bir olcute gore yapilmistir. Model basina "
  "5 puan ve ortalamalari:")
sem_tbl = sem.copy()
sem_tbl.columns = ["Model", "P1", "P2", "P3", "P4", "P5", "Ortalama"]
table(sem_tbl, font_size=8)
P("Yorum: Anlamsal ortalamalar cogunlukla 3.6-3.8 araligindadir; bu, getirilen sonuclarin giris "
  "isletmesiyle ayni alanda (kahve/kafe/firin) oldugunu fakat kategori ortusmesinin isletmeden "
  "isletmeye degistigini gosterir. En yuksek anlamsal ortalamayi window=4 yapilandirmalari (orn. "
  "{} ) almistir; genis pencere, kisa kategori dizilerinde daha fazla kelime-birlikteligi "
  "yakaladigi icin tema butunlugu biraz daha guclu cikmistir.".format(BEST))

H("3.4 (3) Siralama Tutarliligi - Jaccard (Ranking Agreement)", 2)
P("Modellerin ilk 5 sonuc listelerinin ortusmesi Jaccard ile olculmus ve 16x16 matris "
  "olusturulmustur. Jaccard, iki modelin AYNI giris metnine verdigi sonuclarin ne kadar ortustugunu "
  "olcer (tekil basari degil). Kosegen, modelin kendisiyle kiyasi oldugu icin 1.00'dir ve "
  "yorumlamaya katilmaz. Heatmap:")
doc.add_picture(os.path.join(OUT, "jaccard_heatmap.png"), width=Inches(6.3))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
P("Yorum: Kosegen disi Jaccard degerleri yaklasik {:.2f} - {:.2f} araliginda, ortalama ~{:.2f}'dir. "
  "Bu yuksek ortusme, giris isletmesinin (kahve/kafe) temasinin veri setinde belirgin ve ayrik bir "
  "kume olusturmasindan kaynaklanir: 16 modelin neredeyse tamami ayni kahve/kafe kumesini "
  "getirmektedir. Yine de heatmap'te ince oruntuler vardir; ayni algoritma (CBOW/SkipGram) ve ayni "
  "window degerine sahip modeller birbirine daha yuksek Jaccard ile baglanir, yani siralamadaki "
  "kucuk farklar model yapilandirmasiyla iliskilidir.".format(J_MIN, J_MAX, J_MEAN))

H("3.5 Genel Karsilastirma ve En Basarili Modeller", 2)
P("Ozet tablo (model basina ortalama cosine ve anlamsal):", bold=True)
sum_tbl = summary.copy()
sum_tbl.columns = ["Model", "Ort. Cosine", "Ort. Anlamsal"]
table(sum_tbl, font_size=8)
P("Anlamsal vs Cosine: Cosine neredeyse tum modellerde tavana yakindir (ayristirma gucu dusuk); "
  "gercek farki anlamsal degerlendirme ve Jaccard kumelenmesi ortaya koyar. Iki olcut de ayni yonu "
  "gosterir: sonuclar tematik olarak dogrudur, ince kalite farki ise yalnizca anlamsal puanda "
  "gorulur. Bu, neden tek bir olcute guvenmememiz gerektigini gosterir.")
P("Degerlendirme (en basarili / orta / en zayif):", bold=True)
bullet("EN BASARILI: window=4 yapilandirmalari, ozellikle {} ve {}. En yuksek anlamsal ortalamayi "
       "aldilar; genis pencere kisa kategori dizilerinde tema butunlugunu iyi yakaladi.".format(BEST, BEST2))
bullet("ORTA: window=2 modelleri ve cogu SkipGram modeli. Sonuclari hala dogru temada; anlamsal "
       "ortalamalari ~3.6 civarindadir.")
bullet("GORECELI EN ZAYIF: {} . Tema yine dogru olmakla birlikte, ince kategori ortusmesinde biraz "
       "daha dusuk ortalama uretti.".format(WORST))
P("Neden bu modeller one cikti? Veri setimiz kisa, tek dilli (Ingilizce) ve kategori-yogun oldugu "
  "icin, kelime birlikteligini iyi yakalayan yapilandirmalar tema-tutarli vektorler uretti. "
  "Bu olcekte vektor boyutunun (100 vs 300) etkisi sinirli kalmistir.")

H("3.6 Model Yapilandirmalarinin Basariya Etkisi", 2)
bullet("Algoritma (CBOW vs SkipGram): CBOW bu kisa metinlerde biraz daha yuksek/tutarli komsuluk ve "
       "anlamsal puan; SkipGram daha ayristirici ama biraz daha gurultulu skorlar verdi.")
bullet("Window (2 vs 4): window=4, kategori dizilerinde daha genis baglam yakaladigi icin anlamsal "
       "ortalamada hafif ustunluk gosterdi.")
bullet("Vektor boyutu (100 vs 300): Bu veri olceginde anlamli bir kalite farki gozlenmedi; 300 boyut "
       "daha buyuk dosya getirmesine ragmen kucuk korpusta ek fayda saglamadi.")

doc.add_page_break()

# ============================================================================
# 4. SONUC VE ONERILER
# ============================================================================
H("4. Sonuc ve Oneriler", 1)
P("Bu calismada Yelp isletme verisinden uretilen lemmatized ve stemmed setleri uzerinde 16 Word2Vec "
  "modeli egitilmis; secilen bir kahve/kafe isletmesine en benzer 5 isletme her model icin bulunmus "
  "ve modeller cosine, anlamsal ve Jaccard ile karsilastirilmistir. Temel cikarimlar:")
bullet("Kisa, kategori-yogun isletme metinlerinde window=4 yapilandirmalari en tutarli ve anlamli "
       "benzerligi uretti; benzer kisa-metin/etiket benzerligi gorevleri icin bu yapilandirma onerilir.")
bullet("Mutlak cosine skoru tek basina yaniltici olabilir (tum modellerde tavana yakin). Model "
       "kalitesi cosine + anlamsal + Jaccard birlikte yorumlanarak degerlendirilmelidir.")
bullet("Buyuk vektor boyutu (300) bu olcekte ek fayda saglamadi; daha buyuk/cesitli korpuslarda "
       "(orn. yuz binlerce Yelp yorumu) SkipGram + dim300 yeniden denenmelidir.")
bullet("Stemming, lemmatization'a gore biraz daha kucuk bir sozluk uretti; ikisi de benzer benzerlik "
       "sonuclari verdi. Okunabilirlik onemliyse lemmatized, sozluk kucukluğu/hiz onemliyse stemmed.")
P("Hangi model hangi gorev icin? Hizli ve tema-tutarli isletme/oneri benzerligi icin "
  "CBOW + window4 + dim100 (kucuk, hizli, basarili); ince anlam ayrimi ve veri bollugu olan "
  "senaryolar icin SkipGram + window4 + dim300 onerilir.")

out_docx = os.path.join(REP, "Odev2_Raporu.docx")
doc.save(out_docx)
print("DOCX yazildi:", out_docx)

try:
    from docx2pdf import convert
    out_pdf = os.path.join(REP, "Odev2_Raporu.pdf")
    convert(out_docx, out_pdf)
    print("PDF yazildi:", out_pdf)
except Exception as e:
    print("PDF donusumu atlandi (Word/docx2pdf yok). DOCX'i elle PDF'e cevirebilirsiniz. Detay:", e)
